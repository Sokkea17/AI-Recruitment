import io
import re
from typing import Dict, Any, List, Optional
import pypdf
import docx

class DocumentParser:
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            return DocumentParser._extract_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            return DocumentParser._extract_from_docx(file_bytes)
        elif filename_lower.endswith(".txt"):
            return DocumentParser._extract_from_txt(file_bytes)
        elif filename_lower.endswith(".doc"):
            return DocumentParser._extract_from_legacy_doc(file_bytes)
        else:
            raise ValueError(f"Unsupported document format for text extraction: {filename}")

    @staticmethod
    def _extract_from_pdf(file_bytes: bytes) -> str:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n".join(text_parts).strip()

    @staticmethod
    def _extract_from_docx(file_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()

    @staticmethod
    def _extract_from_txt(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="replace").strip()

    @staticmethod
    def _extract_from_legacy_doc(file_bytes: bytes) -> str:
        text = re.sub(rb"[^\x20-\x7E\n\r\t]", b" ", file_bytes).decode("ascii", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 3]
        return "\n".join(lines)

    @staticmethod
    def parse_jd_sections(text: str) -> Dict[str, Any]:
        result = {
            "title": None,
            "department": None,
            "location": None,
            "employment_type": None,
            "salary_range": None,
            "short_description": None,
            "full_description": text,
            "responsibilities": None,
            "requirements": None,
            "education": None,
            "experience": None,
            "skills": None,
            "instructions": None,
            "positions_detected": []
        }

        if not text:
            return result

        lines = [l.strip() for l in text.splitlines() if l.strip()]

        # 1. Extract Title
        title_patterns = [
            r"^(?:job\s+title|position|role)\s*[:\-]\s*(.+)$",
            r"^vacancy\s*[:\-]\s*(.+)$",
            r"^opportunity\s*[:\-]\s*(.+)$"
        ]
        for line in lines[:15]:
            for pat in title_patterns:
                m = re.search(pat, line, re.IGNORECASE)
                if m:
                    result["title"] = m.group(1).strip()
                    break
            if result["title"]:
                break
        
        if not result["title"] and lines:
            first_line = lines[0]
            if len(first_line) < 60 and not any(k in first_line.lower() for k in ["company", "overview", "date", "page", "about"]):
                result["title"] = first_line

        # 2. Extract Department
        dept_match = re.search(r"(?:department|division|team|business\s+unit)\s*[:\-]\s*([^\n]+)", text, re.IGNORECASE)
        if dept_match:
            result["department"] = dept_match.group(1).strip()

        # 3. Extract Location
        loc_match = re.search(r"(?:location|work\s+location|job\s+location|based\s+in)\s*[:\-]\s*([^\n]+)", text, re.IGNORECASE)
        if loc_match:
            result["location"] = loc_match.group(1).strip()

        # 4. Employment Type
        emp_match = re.search(r"(?:employment\s+type|job\s+type|contract\s+type)\s*[:\-]\s*([^\n]+)", text, re.IGNORECASE)
        if emp_match:
            result["employment_type"] = emp_match.group(1).strip()
        else:
            for et in ["Full-time", "Part-time", "Contract", "Internship", "Temporary", "Remote", "Hybrid"]:
                if re.search(r"\b" + re.escape(et) + r"\b", text, re.IGNORECASE):
                    result["employment_type"] = et
                    break

        # 5. Salary
        salary_match = re.search(r"(?:salary|compensation|remuneration|pay\s+range)\s*[:\-]\s*([^\n]+)", text, re.IGNORECASE)
        if salary_match:
            result["salary_range"] = salary_match.group(1).strip()
        else:
            sal_alt = re.search(r"(\$[\d,]+(?:\s*-\s*\$[\d,]+)?(?:\s*(?:per\s+month|per\s+year|/\s*month|/\s*year|monthly|annually))?)", text, re.IGNORECASE)
            if sal_alt and len(sal_alt.group(1)) > 3:
                result["salary_range"] = sal_alt.group(1).strip()

        # Section extraction
        header_definitions = [
            ("summary", [r"job\s+summary", r"role\s+summary", r"about\s+the\s+role", r"position\s+overview", r"job\s+overview", r"about\s+us"]),
            ("responsibilities", [r"responsibilities", r"duties", r"key\s+responsibilities", r"accountabilities", r"what\s+you(?:\'|\s)ll\s+do", r"scope\s+of\s+work"]),
            ("requirements", [r"requirements", r"qualifications", r"key\s+requirements", r"what\s+you\s+need", r"who\s+you\s+are", r"profile"]),
            ("education", [r"education", r"academic\s+qualifications", r"educational\s+background"]),
            ("experience", [r"experience", r"work\s+experience", r"years\s+of\s+experience"]),
            ("skills", [r"skills", r"technical\s+skills", r"competencies", r"core\s+skills"]),
            ("instructions", [r"how\s+to\s+apply", r"application\s+instructions", r"submission\s+details", r"to\s+apply"])
        ]

        extracted_sections = DocumentParser._extract_blocks_by_headers(lines, header_definitions)

        if extracted_sections.get("summary"):
            result["short_description"] = extracted_sections["summary"][:600].strip()
        elif lines:
            desc_candidates = [l for l in lines[1:5] if len(l) > 30]
            if desc_candidates:
                result["short_description"] = " ".join(desc_candidates)[:400]

        if extracted_sections.get("responsibilities"):
            result["responsibilities"] = extracted_sections["responsibilities"].strip()
        if extracted_sections.get("requirements"):
            result["requirements"] = extracted_sections["requirements"].strip()
        if extracted_sections.get("education"):
            result["education"] = extracted_sections["education"].strip()
        if extracted_sections.get("experience"):
            result["experience"] = extracted_sections["experience"].strip()
        if extracted_sections.get("skills"):
            result["skills"] = extracted_sections["skills"].strip()
        if extracted_sections.get("instructions"):
            result["instructions"] = extracted_sections["instructions"].strip()

        # Multi-position detection
        multi_pos = re.findall(r"(?:Position|Job\s+Title|Role)\s*\d*\s*[:\-]\s*([A-Z][A-Za-z0-9\s&/\-]{3,50})", text)
        clean_pos = list(dict.fromkeys([p.strip() for p in multi_pos if len(p.strip()) > 3]))
        if len(clean_pos) > 1:
            result["positions_detected"] = [{"title": p} for p in clean_pos]

        return result

    @staticmethod
    def _extract_blocks_by_headers(lines: List[str], header_definitions: List[tuple]) -> Dict[str, str]:
        results = {}
        current_section = None
        current_content = []

        def match_header(line: str) -> Optional[str]:
            clean = line.strip().lower().rstrip(":.-#")
            for sec_name, patterns in header_definitions:
                for pat in patterns:
                    if re.match(r"^(?:\d+[.\)]\s*)?" + pat + r"\s*$", clean, re.IGNORECASE):
                        return sec_name
            return None

        for line in lines:
            detected_sec = match_header(line)
            if detected_sec:
                if current_section and current_content:
                    results[current_section] = "\n".join(current_content).strip()
                current_section = detected_sec
                current_content = []
            elif current_section:
                current_content.append(line)

        if current_section and current_content:
            results[current_section] = "\n".join(current_content).strip()

        return results

    @staticmethod
    def extract_cv_details(text: str) -> Dict[str, Any]:
        info = {
            "full_name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "experience_snippet": None,
            "education_snippet": None
        }

        email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
        if email_match:
            info["email"] = email_match.group(0).lower()

        phone_match = re.search(r"(?:\+?\d{1,4}[\s.-]?)?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,5}", text)
        if phone_match and len(phone_match.group(0).strip()) >= 8:
            info["phone"] = phone_match.group(0).strip()

        lines = [l.strip() for l in text.splitlines() if l.strip()]
        for line in lines[:8]:
            if line.lower() in ["curriculum vitae", "resume", "cv", "personal information", "contact"]:
                continue
            if info["email"] and info["email"] in line.lower():
                continue
            if info["phone"] and info["phone"] in line:
                continue
            words = line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
                info["full_name"] = line
                break

        if not info["full_name"] and lines:
            first = lines[0]
            if 2 <= len(first.split()) <= 4 and not any(k in first.lower() for k in ["curriculum", "resume", "page", "email"]):
                info["full_name"] = first

        common_keywords = [
            "Python", "FastAPI", "Django", "Flask", "JavaScript", "TypeScript", "React", "Vue", "Node.js",
            "SQL", "PostgreSQL", "MySQL", "MongoDB", "Docker", "Kubernetes", "AWS", "Azure", "GCP",
            "Git", "Linux", "REST API", "GraphQL", "Machine Learning", "Data Analysis", "HR Management",
            "Recruitment", "Talent Acquisition", "Employee Relations", "Payroll", "Communication", "Leadership"
        ]
        found_skills = []
        for kw in common_keywords:
            if re.search(r"\b" + re.escape(kw) + r"\b", text, re.IGNORECASE):
                found_skills.append(kw)
        info["skills"] = found_skills

        return info

document_parser = DocumentParser()
