import os
import docx
from docx.shared import Inches, Pt, RGBColor

out_dir = "/Users/sokkea/telegram-recruitment-bot/sample_jds"
os.makedirs(out_dir, exist_ok=True)

def generate_doc(filename, is_sample=False):
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title
    p = doc.add_paragraph()
    r = p.add_run("JOB DESCRIPTION TEMPLATE" if not is_sample else "JOB DESCRIPTION: HR EXECUTIVE")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(99, 102, 241)
    p.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub_r = sub.add_run("Optimized for automated parsing by the Recruitment Assistant" if not is_sample else "Sample recruitment document")
    sub_r.font.size = Pt(10)
    sub_r.font.italic = True
    sub_r.font.color.rgb = RGBColor(100, 116, 139)
    sub.paragraph_format.space_after = Pt(14)

    # Metadata Table
    table = doc.add_table(rows=5, cols=2)
    meta = [
        ("Position Title:", "[e.g. Senior Legal Executive]" if not is_sample else "HR Executive"),
        ("Department:", "[e.g. Legal & Compliance]" if not is_sample else "Human Resources & Admin"),
        ("Location:", "[e.g. Phnom Penh, Cambodia]" if not is_sample else "Phnom Penh, Cambodia"),
        ("Employment Type:", "[Full-time / Part-time / Contract]" if not is_sample else "Full-time"),
        ("Salary Range:", "[e.g. $1,200 - $1,800]" if not is_sample else "$600 - $900 per month")
    ]
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(k).font.bold = True
        row.cells[1].paragraphs[0].add_run(v)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def section(title, items):
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.font.bold = True
        hr.font.size = Pt(12.5)
        hr.font.color.rgb = RGBColor(30, 41, 59)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

        if isinstance(items, list):
            for it in items:
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(it).font.size = Pt(10.5)
                bp.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            p.add_run(items).font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(6)

    if not is_sample:
        section("Job Summary", "Provide a brief 2-3 sentence overview of this role. This summary will be shown to candidates browsing jobs on Telegram.")
        section("Key Responsibilities", [
            "Lead and execute department operations and assigned projects.",
            "Collaborate with internal teams and external partners.",
            "Prepare operational reports and performance tracking metrics.",
            "Ensure full adherence to company standards and regulatory guidelines."
        ])
        section("Requirements & Qualifications", [
            "Bachelor's degree in a relevant field.",
            "Minimum 2-3 years of proven experience in a similar position.",
            "Strong analytical and problem-solving skills.",
            "Effective communication skills in English."
        ])
        section("Education", "Bachelor's degree or higher in relevant discipline.")
        section("Experience", "Minimum 2-4 years of industry experience.")
        section("Skills", "Communication, Project Management, Critical Thinking, Relevant Technical Tools.")
        section("How to Apply", "Submit your updated CV in PDF or Word format through our Telegram Recruitment Bot.")
    else:
        section("Job Summary", "We are seeking an HR Executive to coordinate end-to-end recruitment, maintain employee records, and support daily HR operations.")
        section("Key Responsibilities", [
            "Manage recruitment cycles: source candidates, screen CVs, and schedule interviews.",
            "Maintain confidential employee profiles, contracts, and attendance records.",
            "Assist with monthly payroll preparation and employee benefits.",
            "Coordinate new employee onboarding and induction sessions."
        ])
        section("Requirements & Qualifications", [
            "Bachelor's degree in Human Resources, Business, or related field.",
            "Minimum 2 years experience in HR operations or talent acquisition.",
            "Sound knowledge of labor regulations and HR practices.",
            "Proficiency in Microsoft Office (Word, Excel) and English fluency."
        ])
        section("Education", "Bachelor's degree in Human Resource Management or Business Administration.")
        section("Experience", "2+ years in recruitment or HR administration.")
        section("Skills", "Recruitment, Sourcing, Interviewing, Labor Law, MS Excel, Communication.")
        section("How to Apply", "Submit your CV in PDF or DOCX via our Telegram Recruitment Bot.")

    doc.save(os.path.join(out_dir, filename))
    print(f"Generated: {filename}")

generate_doc("Job_Description_Template.docx", is_sample=False)
generate_doc("Sample_HR_Executive_JD.docx", is_sample=True)
